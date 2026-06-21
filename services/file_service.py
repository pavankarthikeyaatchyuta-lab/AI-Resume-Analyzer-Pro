import hashlib
from io import BytesIO

import docx
import pdfplumber
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from reportlab.pdfgen import canvas

from .analysis_service import RESUME_HEADINGS


def extract_text(uploaded_file):
    uploaded_file.seek(0)
    text = ""

    if uploaded_file.name.lower().endswith(".pdf"):
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif uploaded_file.name.lower().endswith(".docx"):
        document = docx.Document(uploaded_file)
        for para in document.paragraphs:
            text += para.text + "\n"

    uploaded_file.seek(0)
    return text.strip()


def get_uploaded_file_signature(uploaded_file):
    uploaded_file.seek(0)
    content = uploaded_file.getvalue()
    uploaded_file.seek(0)
    digest = hashlib.sha256(content).hexdigest()
    return f"{uploaded_file.name}:{len(content)}:{digest}"


def clear_generated_outputs(st):
    st.session_state["analysis"] = None
    st.session_state["analysis_pdf"] = None
    st.session_state["improved_resume"] = None
    st.session_state["improved_resume_docx"] = None
    st.session_state["improved_resume_pdf"] = None
    st.session_state["scores"] = None


def generate_pdf_bytes(report_text):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    page_width, page_height = letter
    left_margin = 40
    right_margin = 40
    top_margin = 42
    bottom_margin = 42
    font_name = "Helvetica"
    font_size = 10
    line_height = 14
    printable_width = page_width - left_margin - right_margin

    def start_page():
        pdf.setFont(font_name, font_size)
        return page_height - top_margin

    y_position = start_page()

    for raw_line in report_text.splitlines():
        wrapped_lines = simpleSplit(raw_line, font_name, font_size, printable_width) or [""]
        for line in wrapped_lines:
            if y_position <= bottom_margin:
                pdf.showPage()
                y_position = start_page()
            pdf.drawString(left_margin, y_position, line)
            y_position -= line_height

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def generate_resume_docx_bytes(text):
    document = docx.Document()
    section = document.sections[0]
    section.top_margin = Pt(42)
    section.bottom_margin = Pt(42)
    section.left_margin = Pt(42)
    section.right_margin = Pt(42)

    base_style = document.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(11)

    lines = [line.strip() for line in text.splitlines()]
    for line in lines:
        if not line:
            continue

        if line.upper() in RESUME_HEADINGS:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(12)
            continue

        if line.startswith("- "):
            para = document.add_paragraph(style="List Bullet")
            para.add_run(line[2:].strip())
            continue

        if not document.paragraphs:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(15)
            continue

        para = document.add_paragraph(line)
        para.paragraph_format.space_after = Pt(4)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
