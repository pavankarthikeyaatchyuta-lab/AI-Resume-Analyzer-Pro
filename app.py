import hashlib
import json
import os
import re
import time
from io import BytesIO

import docx
import pdfplumber
import plotly.graph_objects as go
import streamlit as st
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from groq import APIConnectionError, APIStatusError, APITimeoutError, AuthenticationError, BadRequestError, Groq, RateLimitError
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from reportlab.pdfgen import canvas
from streamlit.errors import StreamlitSecretNotFoundError

from frontend import clean_section_text, empty_state, render_background_fx, render_hero, render_metrics, render_report_block, render_styles


st.set_page_config(page_title="AI Resume Analyzer Pro", page_icon="A", layout="wide")


RESUME_HEADINGS = {
    "PROFILE SUMMARY",
    "PROFESSIONAL SUMMARY",
    "SUMMARY",
    "EDUCATION",
    "TECHNICAL SKILLS",
    "SKILLS",
    "PROJECTS",
    "EXPERIENCE",
    "WORK EXPERIENCE",
    "CERTIFICATIONS & TRAINING",
    "CERTIFICATIONS",
    "ACHIEVEMENTS & ACTIVITIES",
    "ACHIEVEMENTS",
    "KEY STRENGTHS",
    "STRENGTHS",
    "CONTACT",
}

ANALYSIS_SECTION_ALIASES = {
    "top_suggestions": ["Top 5 Suggestions", "Top Suggestions"],
    "missing_keywords": ["Missing Keywords"],
    "skill_gap_analysis": ["Skill Gap Analysis"],
    "improved_bullets": ["Improved Professional Bullet Points", "Improved Bullet Points"],
}

GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_FALLBACK_MODEL = os.getenv("GROQ_FALLBACK_MODEL", "llama-3.1-8b-instant")
GROQ_REQUEST_TIMEOUT_SECONDS = float(os.getenv("GROQ_REQUEST_TIMEOUT_SECONDS", "40"))
GROQ_MAX_ATTEMPTS = 3
GROQ_CALL_COOLDOWN_SECONDS = int(os.getenv("GROQ_CALL_COOLDOWN_SECONDS", "8"))

STOP_WORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "in",
    "into",
    "is",
    "it",
    "of",
    "on",
    "or",
    "that",
    "the",
    "this",
    "to",
    "with",
    "you",
    "your",
    "we",
    "our",
    "will",
    "should",
    "must",
    "have",
    "has",
    "had",
    "job",
    "role",
    "description",
    "requirements",
    "responsibilities",
    "preferred",
    "qualification",
    "qualifications",
}


def init_session_state():
    # Keep all generated artifacts in session state so tab switches do not lose results.
    defaults = {
        "analysis": None,
        "analysis_pdf": None,
        "improved_resume": None,
        "improved_resume_docx": None,
        "improved_resume_pdf": None,
        "scores": None,
        "last_uploaded_signature": None,
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


def reset_app():
    for key in [
        "analysis",
        "analysis_pdf",
        "improved_resume",
        "improved_resume_docx",
        "improved_resume_pdf",
        "scores",
        "last_uploaded_signature",
        "resume_uploader",
        "job_role_input",
        "chart_selector",
    ]:
        if key in st.session_state:
            del st.session_state[key]
    st.rerun()


def get_groq_client():
    try:
        api_key = st.secrets.get("GROQ_API_KEY")
    except StreamlitSecretNotFoundError:
        api_key = None

    if not api_key:
        api_key = os.getenv("GROQ_API_KEY")

    if not api_key:
        st.error("Missing Groq API key. Add `GROQ_API_KEY` to Streamlit secrets or environment variables.")
        st.stop()
    return Groq(api_key=api_key, timeout=GROQ_REQUEST_TIMEOUT_SECONDS, max_retries=0)


def extract_text(uploaded_file):
    uploaded_file.seek(0)
    text = ""

    if uploaded_file.name.lower().endswith(".pdf"):
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif uploaded_file.name.lower().endswith(".docx"):
        document = docx.Document(uploaded_file)
        for para in document.paragraphs:
            text += para.text + "\n"

    uploaded_file.seek(0)
    return text.strip()


def get_uploaded_file_signature(uploaded_file):
    # Hash the full file content so replacing a resume with the same filename still resets stale results.
    uploaded_file.seek(0)
    content = uploaded_file.getvalue()
    uploaded_file.seek(0)
    digest = hashlib.sha256(content).hexdigest()
    return f"{uploaded_file.name}:{len(content)}:{digest}"


def tokenize_for_overlap(text):
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9+#.-]*", text.lower())
    return {token for token in tokens if len(token) > 2 and token not in STOP_WORDS and not token.isdigit()}


def clear_generated_outputs():
    st.session_state["analysis"] = None
    st.session_state["analysis_pdf"] = None
    st.session_state["improved_resume"] = None
    st.session_state["improved_resume_docx"] = None
    st.session_state["improved_resume_pdf"] = None
    st.session_state["scores"] = None


def get_groq_model_candidates():
    candidates = [GROQ_MODEL]
    if GROQ_FALLBACK_MODEL and GROQ_FALLBACK_MODEL != GROQ_MODEL:
        candidates.append(GROQ_FALLBACK_MODEL)
    return candidates


def is_model_unavailable_error(exception):
    if isinstance(exception, APIStatusError) and getattr(exception, "status_code", None) in {400, 404}:
        message = str(exception).lower()
        return "model" in message and any(token in message for token in ["not found", "deprecated", "unavailable", "does not exist"])
    return False


def is_transient_groq_error(exception):
    if isinstance(exception, (RateLimitError, APITimeoutError, APIConnectionError)):
        return True
    if isinstance(exception, APIStatusError):
        return getattr(exception, "status_code", 0) >= 500
    return False


def classify_groq_error(exception):
    if isinstance(exception, AuthenticationError):
        return "auth"
    if isinstance(exception, BadRequestError):
        return "bad_request"
    if isinstance(exception, RateLimitError):
        return "rate_limit"
    if is_model_unavailable_error(exception):
        return "model_unavailable"
    if isinstance(exception, APIStatusError) and getattr(exception, "status_code", 0) < 500:
        return "bad_request"
    if is_transient_groq_error(exception):
        return "transient"
    return "generic"


def format_groq_error_message(exception):
    category = classify_groq_error(exception)
    if category == "rate_limit":
        return "Too many requests right now, please wait a moment and try again."
    if category == "auth":
        return "API key issue — contact the app owner."
    return "Something went wrong, please try again."


def call_groq_with_retry(client, *, messages, temperature=0, response_format=None):
    last_exception = None
    for model_index, model_name in enumerate(get_groq_model_candidates()):
        for attempt in range(1, GROQ_MAX_ATTEMPTS + 1):
            request_kwargs = {
                "model": model_name,
                "temperature": temperature,
                "messages": messages,
            }
            if response_format is not None:
                request_kwargs["response_format"] = response_format

            try:
                return client.chat.completions.create(**request_kwargs)
            except Exception as exception:
                last_exception = exception
                category = classify_groq_error(exception)

                if category in {"auth", "bad_request"}:
                    raise

                if category == "model_unavailable":
                    break

                if attempt < GROQ_MAX_ATTEMPTS and category in {"transient", "rate_limit"}:
                    time.sleep(2 ** (attempt - 1))
                    continue

                if attempt < GROQ_MAX_ATTEMPTS and category == "generic":
                    time.sleep(2 ** (attempt - 1))
                    continue

                break

        if last_exception is not None and classify_groq_error(last_exception) in {"auth", "bad_request"}:
            raise last_exception

        if model_index < len(get_groq_model_candidates()) - 1 and classify_groq_error(last_exception) in {"transient", "rate_limit", "model_unavailable", "generic"}:
            continue

    raise last_exception


def groq_call_allowed(action_key, cooldown_seconds=GROQ_CALL_COOLDOWN_SECONDS):
    now = time.time()
    last_call_time = st.session_state.get(f"last_groq_call_at_{action_key}")
    if last_call_time is not None and now - last_call_time < cooldown_seconds:
        remaining_seconds = max(1, int(round(cooldown_seconds - (now - last_call_time))))
        st.warning(f"Please wait a few seconds before trying again. ({remaining_seconds}s remaining)")
        return False

    st.session_state[f"last_groq_call_at_{action_key}"] = now
    return True


def build_job_description_context(job_role, resume_text):
    word_count = len(re.findall(r"\b\w+\b", job_role or ""))
    if word_count < 8:
        return ""

    job_tokens = tokenize_for_overlap(job_role)
    resume_tokens = tokenize_for_overlap(resume_text)
    if not job_tokens:
        return ""

    overlap = len(job_tokens & resume_tokens) / max(len(job_tokens | resume_tokens), 1)
    missing_keywords = sorted(job_tokens - resume_tokens)
    highlighted_missing = ", ".join(missing_keywords[:20])

    if highlighted_missing:
        return (
            "Job description grounding:\n"
            f"- Overlap similarity: {overlap:.2f}\n"
            f"- Keywords present in the job description but not the resume: {highlighted_missing}"
        )

    return f"Job description grounding:\n- Overlap similarity: {overlap:.2f}"


def normalize_analysis_items(items, limit=None):
    if isinstance(items, str):
        items = re.split(r"\s*,\s*|\s*\n\s*", items)
    elif not isinstance(items, (list, tuple)):
        items = []

    normalized_items = []
    seen = set()

    for item in items or []:
        cleaned_item = re.sub(r"\s+", " ", str(item)).strip(" \t\r\n-•")
        if not cleaned_item:
            continue

        lowered = cleaned_item.lower()
        if lowered in seen:
            continue

        seen.add(lowered)
        normalized_items.append(cleaned_item)
        if limit is not None and len(normalized_items) >= limit:
            break

    return normalized_items


def analysis_json_to_legacy_format(parsed_json) -> dict:
    resume_score = int(parsed_json.get("resume_score", 70) or 0)
    ats_score = int(parsed_json.get("ats_score", 65) or 0)
    job_suitability_score = int(parsed_json.get("job_suitability_score", 60) or 0)

    scores = {
        "Resume Score": max(0, min(100, resume_score)),
        "ATS Score": max(0, min(100, ats_score)),
        "Job Suitability": max(0, min(100, job_suitability_score)),
    }

    top_suggestions = normalize_analysis_items(parsed_json.get("top_suggestions"), limit=5)
    missing_keywords = normalize_analysis_items(parsed_json.get("missing_keywords"), limit=10)
    skill_gap_analysis = normalize_analysis_items(parsed_json.get("skill_gap_analysis"), limit=2)
    improved_bullets = normalize_analysis_items(parsed_json.get("improved_bullets"), limit=3)

    analysis_sections = [
        f"Resume Score: {scores['Resume Score']}",
        f"ATS Score: {scores['ATS Score']}",
        f"Job Suitability Score: {scores['Job Suitability']}",
        "",
        "Top 5 Suggestions",
        "\n".join(f"- {item}" for item in top_suggestions),
        "",
        "Missing Keywords",
        ", ".join(missing_keywords),
        "",
        "Skill Gap Analysis",
        "\n".join(f"- {item}" for item in skill_gap_analysis),
        "",
        "Improved Professional Bullet Points",
        "\n".join(f"- {item}" for item in improved_bullets),
    ]

    analysis_text = "\n".join(part for part in analysis_sections if part is not None).strip()

    return {
        "analysis": analysis_text,
        "scores": scores,
    }


def analyze_resume(client, text, job_role):
    # The response format is tightly constrained because the analysis parser depends on stable section headings.
    job_context = build_job_description_context(job_role, text)
    prompt = f"""
You are a senior ATS recruiter.

Analyze this resume for the job role: {job_role}

Return a single JSON object with this exact schema:
{{
  "resume_score": 0,
  "ats_score": 0,
  "job_suitability_score": 0,
  "top_suggestions": ["...", "...", "...", "...", "..."],
  "missing_keywords": ["...", "..."],
  "skill_gap_analysis": ["...", "..."],
  "improved_bullets": ["...", "...", "..."]
}}

Rules:
- Return JSON only.
- Do not wrap the response in markdown fences.
- Use integers between 0 and 100 for the scores.
- Keep top_suggestions at exactly 5 items.
- Keep missing_keywords between 6 and 10 items.
- Keep skill_gap_analysis at exactly 2 items.
- Keep improved_bullets at exactly 3 items.
- Use short strings in each array item.

{job_context}

Resume:
{text}
"""

    completion = call_groq_with_retry(
        client,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        response_format={"type": "json_object"},
    )
    raw_response = completion.choices[0].message.content

    try:
        parsed_json = json.loads(raw_response)
        if isinstance(parsed_json, dict):
            return analysis_json_to_legacy_format(parsed_json)
    except (TypeError, json.JSONDecodeError):
        pass

    legacy_text = raw_response.strip()
    return {
        "analysis": legacy_text,
        "scores": parse_scores(legacy_text),
    }


def generate_updated_resume(client, text):
    prompt = f"""
You are an expert ATS resume writer.

Rewrite the resume as a polished final resume, but keep all factual information intact.

Rules:
- Return ONLY the final resume content.
- Do NOT include explanations, introductions, notes, or "here is your resume".
- Do NOT include "I made the following changes" or any commentary.
- Preserve contact details, education, projects, certifications, strengths, and skills.
- Use professional section headings.
- Use concise bullet points where appropriate.
- Output plain text only.

Original Resume:
{text}
"""

    completion = call_groq_with_retry(
        client,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    return completion.choices[0].message.content


def strip_markdown_links(text):
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)


def clean_resume_text(text):
    text = strip_markdown_links(text.replace("\r\n", "\n")).strip()
    text = re.sub(r"^Here(?:'s| is).*$", "", text, flags=re.IGNORECASE | re.MULTILINE)
    text = re.sub(r"^I made the following changes.*$", "", text, flags=re.IGNORECASE | re.MULTILINE)
    text = re.sub(r"\n\s*\d+\.\s+\*\*.*$", "", text, flags=re.MULTILINE)
    text = text.replace("**", "")
    text = re.sub(r"^[ \t]*[\*\-]\s+", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"^[ \t]*\+\s+", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"[ \t]+", " ", text)

    lines = [line.strip() for line in text.splitlines()]
    cleaned = []
    for line in lines:
        if not line:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue

        lower = line.lower()
        if lower.startswith("here's a rewritten version") or lower.startswith("i made the following changes"):
            continue
        if re.match(r"^\d+\.\s+", line) and "improved" in lower and "resume" not in lower:
            continue
        cleaned.append(line)

    cleaned_text = "\n".join(cleaned).strip()
    cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text)
    return cleaned_text


def split_analysis_sections(analysis):
    sections = {
        "top_suggestions": "",
        "missing_keywords": "",
        "skill_gap_analysis": "",
        "improved_bullets": "",
        "other_notes": "",
    }

    score_patterns = [
        r"Resume Score:\s*\d{1,3}",
        r"ATS Score:\s*\d{1,3}",
        r"Job Suitability Score:\s*\d{1,3}",
    ]
    # Strip the score header first so we only parse the narrative sections below it.
    working_text = analysis
    for pattern in score_patterns:
        working_text = re.sub(pattern, "", working_text, flags=re.IGNORECASE)

    heading_matches = []
    for key, aliases in ANALYSIS_SECTION_ALIASES.items():
        alias_pattern = "|".join(re.escape(alias) for alias in aliases)
        pattern = rf"(?im)^\s*(?:{alias_pattern})\s*:?\s*"
        for match in re.finditer(pattern, working_text):
            heading_matches.append((match.start(), match.end(), key))

    heading_matches.sort(key=lambda item: item[0])

    consumed_ranges = []
    # Slice each heading block until the next known heading so every section stays isolated in the UI.
    for index, (start, end, key) in enumerate(heading_matches):
        next_start = heading_matches[index + 1][0] if index + 1 < len(heading_matches) else len(working_text)
        section_content = working_text[end:next_start]
        cleaned_content = clean_section_text(section_content)
        if cleaned_content and not sections[key]:
            sections[key] = cleaned_content
            consumed_ranges.append((start, next_start))

    other_notes_parts = []
    cursor = 0
    # Anything outside the known headings is preserved as "Additional Notes" instead of being discarded.
    for start, end in consumed_ranges:
        leftover = clean_section_text(working_text[cursor:start])
        if leftover:
            other_notes_parts.append(leftover)
        cursor = end

    trailing_leftover = clean_section_text(working_text[cursor:])
    if trailing_leftover:
        other_notes_parts.append(trailing_leftover)

    sections["other_notes"] = "\n\n".join(part for part in other_notes_parts if part)
    return sections


def parse_scores(result):
    patterns = {
        "Resume Score": r"Resume Score:\s*(\d{1,3})",
        "ATS Score": r"ATS Score:\s*(\d{1,3})",
        "Job Suitability": r"Job Suitability Score:\s*(\d{1,3})",
    }
    defaults = {"Resume Score": 70, "ATS Score": 65, "Job Suitability": 60}

    scores = {}
    for label, pattern in patterns.items():
        match = re.search(pattern, result, flags=re.IGNORECASE)
        value = int(match.group(1)) if match else defaults[label]
        scores[label] = max(0, min(100, value))
    return scores


def extract_keywords_from_analysis(analysis):
    match = re.search(
        r"Missing Keywords\s*:?\s*(.*?)(?:\n[A-Z][A-Za-z ]+\n|\nSkill Gap Analysis|\Z)",
        analysis,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if not match:
        return []
    content = match.group(1)
    items = re.split(r"[\n,•\-]+", content)
    keywords = [item.strip(" :.") for item in items if item.strip(" :.")]
    return keywords[:5]


def create_score_chart(scores, chart_type):
    labels = list(scores.keys())
    values = list(scores.values())
    colors = ["#67e8f9", "#8b5cf6", "#f59e0b"]

    if chart_type == "Vertical Bar":
        fig = go.Figure([go.Bar(x=labels, y=values, marker=dict(color=colors), text=values, textposition="outside")])
        fig.update_yaxes(range=[0, 100], gridcolor="rgba(148,163,184,0.16)")
    elif chart_type == "Horizontal Bar":
        fig = go.Figure(
            [go.Bar(y=labels, x=values, orientation="h", marker=dict(color=colors), text=values, textposition="outside")]
        )
        fig.update_xaxes(range=[0, 100], gridcolor="rgba(148,163,184,0.16)")
    elif chart_type == "Pie Chart":
        fig = go.Figure([go.Pie(labels=labels, values=values, marker=dict(colors=colors), hole=0.52)])
    else:
        fig = go.Figure(
            [
                go.Scatterpolar(
                    r=values + values[:1],
                    theta=labels + labels[:1],
                    fill="toself",
                    fillcolor="rgba(103, 232, 249, 0.18)",
                    line=dict(color="#67e8f9", width=3),
                )
            ]
        )
        fig.update_polars(radialaxis=dict(range=[0, 100], visible=True, gridcolor="rgba(148,163,184,0.16)"))

    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#e5eefc"),
        margin=dict(l=20, r=20, t=25, b=20),
        showlegend=False,
    )
    return fig


def create_gauge_chart(scores):
    fig = go.Figure()
    domains = [(0.00, 0.30), (0.35, 0.65), (0.70, 1.00)]
    colors = ["#67e8f9", "#8b5cf6", "#f59e0b"]

    for index, (label, value) in enumerate(scores.items()):
        fig.add_trace(
            go.Indicator(
                mode="gauge+number",
                value=value,
                title={"text": label, "font": {"size": 17, "color": "#e5eefc"}},
                domain={"x": [domains[index][0], domains[index][1]], "y": [0, 1]},
                gauge={
                    "axis": {"range": [0, 100], "tickcolor": "#94a3b8"},
                    "bar": {"color": colors[index]},
                    "bgcolor": "rgba(15,23,42,0.92)",
                    "borderwidth": 1,
                    "bordercolor": "rgba(148,163,184,0.24)",
                },
            )
        )

    fig.update_layout(
        height=320,
        margin=dict(l=20, r=20, t=20, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#e5eefc"),
    )
    return fig


def generate_pdf_bytes(report_text):
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    page_width, page_height = letter
    left_margin = 40
    right_margin = 40
    top_margin = 42
    bottom_margin = 42
    font_name = "Helvetica"
    font_size = 10
    line_height = 14
    printable_width = page_width - left_margin - right_margin

    def start_page():
        pdf.setFont(font_name, font_size)
        return page_height - top_margin

    y_position = start_page()

    # Wrap long lines and paginate manually so large analyses do not get cut off in the exported PDF.
    for raw_line in report_text.splitlines():
        wrapped_lines = simpleSplit(raw_line, font_name, font_size, printable_width) or [""]
        for line in wrapped_lines:
            if y_position <= bottom_margin:
                pdf.showPage()
                y_position = start_page()
            pdf.drawString(left_margin, y_position, line)
            y_position -= line_height

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def generate_resume_docx_bytes(text):
    document = docx.Document()
    section = document.sections[0]
    section.top_margin = Pt(42)
    section.bottom_margin = Pt(42)
    section.left_margin = Pt(42)
    section.right_margin = Pt(42)

    base_style = document.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(11)

    # Use lightweight text heuristics to turn plain-text LLM output into a readable DOCX structure.
    lines = [line.strip() for line in text.splitlines()]
    for line in lines:
        if not line:
            continue

        if line.upper() in RESUME_HEADINGS:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(12)
            continue

        if line.startswith("- "):
            para = document.add_paragraph(style="List Bullet")
            para.add_run(line[2:].strip())
            continue

        if not document.paragraphs:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(15)
            continue

        para = document.add_paragraph(line)
        para.paragraph_format.space_after = Pt(4)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


init_session_state()
render_styles()
render_background_fx()
client = get_groq_client()

current_upload_signature = st.session_state.get("last_uploaded_signature")

with st.sidebar:
    st.markdown("## Control Deck")
    st.caption("Upload a resume, set the role, then launch the scan.")
    resume_file = st.file_uploader("Resume File", type=["pdf", "docx"], key="resume_uploader")
    job_role = st.text_input("Target Job Role", key="job_role_input", placeholder="Machine Learning Engineer")
    chart_type = st.selectbox(
        "Visualization Type",
        ["Vertical Bar", "Horizontal Bar", "Pie Chart", "Radar Chart", "Gauge Dashboard"],
        key="chart_selector",
    )
    analyze_btn = st.button("Analyze Resume", type="primary", use_container_width=True)
    reset_btn = st.button("Reset", use_container_width=True)

if reset_btn:
    reset_app()

if resume_file:
    current_signature = get_uploaded_file_signature(resume_file)
    if current_signature != current_upload_signature:
        # Clear generated content as soon as a different file is uploaded, even if the name did not change.
        clear_generated_outputs()
        st.session_state["last_uploaded_signature"] = current_signature

if analyze_btn:
    if not resume_file:
        st.warning("Please upload a resume.")
    elif not groq_call_allowed("shared"):
        pass
    else:
        try:
            with st.spinner("Analyzing resume..."):
                resume_text = extract_text(resume_file)
                if not resume_text:
                    st.warning("We couldn't extract any text from that file. Please upload a text-based PDF or DOCX.")
                else:
                    analysis_result = analyze_resume(client, resume_text, job_role or "General role")
                    st.session_state["analysis"] = analysis_result["analysis"]
                    st.session_state["scores"] = analysis_result["scores"]
                    st.session_state["analysis_pdf"] = generate_pdf_bytes(analysis_result["analysis"])
                    st.session_state["improved_resume"] = None
                    st.session_state["improved_resume_docx"] = None
                    st.session_state["improved_resume_pdf"] = None
        except Exception as exc:
            clear_generated_outputs()
            st.error(format_groq_error_message(exc))

render_hero(st.session_state["scores"])
render_metrics(st.session_state["scores"])

tab1, tab2, tab3 = st.tabs(["Analysis", "Visuals", "Improved Resume"])

with tab1:
    st.subheader("Resume Analysis")
    if st.session_state["analysis"]:
        sections = split_analysis_sections(st.session_state["analysis"])
        render_report_block("Top Suggestions", sections["top_suggestions"], variant="list")
        render_report_block("Missing Keywords", sections["missing_keywords"], variant="keywords")
        render_report_block("Skill Gap Analysis", sections["skill_gap_analysis"], variant="insights")
        render_report_block("Improved Bullet Points", sections["improved_bullets"], variant="list")

        if sections["other_notes"]:
            render_report_block("Additional Notes", sections["other_notes"], variant="insights")

        st.download_button(
            "Download Analysis Report",
            st.session_state["analysis_pdf"],
            file_name="resume_analysis_report.pdf",
            mime="application/pdf",
        )
    else:
        empty_state("Run the analysis to see the recruiter-style report here.")

with tab2:
    st.subheader("Visual Analytics")
    if st.session_state["scores"]:
        if chart_type == "Gauge Dashboard":
            figure = create_gauge_chart(st.session_state["scores"])
        else:
            figure = create_score_chart(st.session_state["scores"], chart_type)
        st.plotly_chart(figure, use_container_width=True)
    else:
        empty_state("Charts will appear here after analysis.")

with tab3:
    st.subheader("Improved Resume")
    generate_improved = st.button("Generate Improved Resume")

    if generate_improved:
        if not resume_file:
            st.warning("Please upload a resume first.")
        elif not groq_call_allowed("shared"):
            pass
        else:
            try:
                with st.spinner("Generating improved resume..."):
                    resume_text = extract_text(resume_file)
                    if not resume_text:
                        st.warning("We couldn't extract any text from that file. Please upload a text-based PDF or DOCX.")
                    else:
                        improved_resume_raw = generate_updated_resume(client, resume_text)
                        improved_resume = clean_resume_text(improved_resume_raw)
                        st.session_state["improved_resume"] = improved_resume
                        st.session_state["improved_resume_docx"] = generate_resume_docx_bytes(improved_resume)
                        st.session_state["improved_resume_pdf"] = generate_pdf_bytes(improved_resume)
            except Exception as exc:
                st.session_state["improved_resume"] = None
                st.session_state["improved_resume_docx"] = None
                st.session_state["improved_resume_pdf"] = None
                st.error(format_groq_error_message(exc))

    if st.session_state["improved_resume"]:
        st.text_area("Clean Resume Output", st.session_state["improved_resume"], height=520)
        download_col1, download_col2, download_col3 = st.columns(3)
        with download_col1:
            st.download_button(
                "Download PDF",
                st.session_state["improved_resume_pdf"],
                file_name="improved_resume.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        with download_col2:
            st.download_button(
                "Download DOCX",
                st.session_state["improved_resume_docx"],
                file_name="improved_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with download_col3:
            st.download_button(
                "Download TXT",
                st.session_state["improved_resume"],
                file_name="improved_resume.txt",
                mime="text/plain",
                use_container_width=True,
            )
    else:
        empty_state("Generate an improved version of the uploaded resume here.")
